import os
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    """Sets the background color of a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Sets the internal margins (padding) of a cell in dxa (1/20 of a point)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def apply_text_formatting(run, font_name="Times New Roman", size_pt=12, bold=False, italic=False):
    """Applies font family, size, and styles to a text run."""
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic

def add_runs_to_paragraph(p, text):
    """Parses text for markdown bold (**text**) and italic (*text*) and adds runs to paragraph."""
    # Split by bold first
    bold_parts = re.split(r'(\*\*.*?\*\*)', text)
    for bp in bold_parts:
        is_bold = False
        content = bp
        if bp.startswith('**') and bp.endswith('**'):
            is_bold = True
            content = bp[2:-2]
        
        # Now split content by italic
        italic_parts = re.split(r'(\*.*?\*)', content)
        for ip in italic_parts:
            is_italic = False
            sub_content = ip
            if ip.startswith('*') and ip.endswith('*'):
                is_italic = True
                sub_content = ip[1:-1]
            
            if sub_content:
                run = p.add_run(sub_content)
                apply_text_formatting(run, bold=is_bold, italic=is_italic)

def process_paragraph_block(doc, text, in_references=False):
    text_clean = text.strip()
    if not text_clean:
        return
        
    # Check Heading 1
    is_fully_bold = text_clean.startswith("**") and text_clean.endswith("**")
    inner_text = text_clean[2:-2].strip() if is_fully_bold else text_clean
    
    h1_keywords = ["abstract", "references", "table of contents", "list of tables", "list of figures", "list of acronyms/abbreviations"]
    if is_fully_bold and (inner_text.lower() in h1_keywords or inner_text.upper().startswith("CHAPTER ")):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT if (inner_text.upper().startswith("CHAPTER") or inner_text.upper() == "REFERENCES") else WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(24)
        p.paragraph_format.space_after = Pt(18)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(inner_text)
        apply_text_formatting(run, size_pt=16, bold=True)
        return
        
    # Check Heading 2 (X.Y)
    h2_match = re.match(r'^\*\*(\d+\.\d+)\s+(.*?)\*\*$', text_clean)
    if h2_match:
        num, rest = h2_match.groups()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(12)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(f"{num} {rest}")
        apply_text_formatting(run, size_pt=14, bold=True)
        return

    # Check Heading 3 (X.Y.Z)
    h3_match = re.match(r'^\*\*(\d+\.\d+\.\d+)\s+(.*?)\*\*$', text_clean)
    if h3_match:
        num, rest = h3_match.groups()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(f"{num} {rest}")
        apply_text_formatting(run, size_pt=12, bold=True)
        return

    # Check Table and Figure Captions
    is_figure_caption = re.match(r'^\*\*Figure\s+\d+(\.\d+)?:.*?\*\*$', text_clean) or re.match(r'^Figure\s+\d+(\.\d+)?:.*', text_clean)
    is_table_caption = re.match(r'^\*\*Table\s+\d+(\.\d+)?:.*?\*\*$', text_clean) or re.match(r'^Table\s+\d+(\.\d+)?:.*', text_clean)
    
    if is_figure_caption or is_table_caption:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(12)
        p.paragraph_format.keep_with_next = True
        
        caption_text = text_clean[2:-2].strip() if text_clean.startswith("**") and text_clean.endswith("**") else text_clean
        
        run = p.add_run(caption_text)
        apply_text_formatting(run, size_pt=10, bold=True, italic=True)
        return

    # Normal paragraph
    p = doc.add_paragraph()
    if in_references:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(12)
    add_runs_to_paragraph(p, text_clean)

def process_bullet_list_block(doc, lines):
    text = " ".join([line.strip() for line in lines if line.strip()])
    if not text:
        return
        
    if text.startswith("- "):
        text = text[2:].strip()
    elif text.startswith("* "):
        text = text[2:].strip()
        
    p = doc.add_paragraph(style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    
    add_runs_to_paragraph(p, text)

def process_numbered_list_block(doc, lines):
    text = " ".join([line.strip() for line in lines if line.strip()])
    if not text:
        return
        
    num = ""
    rest = text
    
    m1 = re.match(r'^\*\*(\d+)\.\*\*\s*(.*)', text)
    m2 = re.match(r'^(\d+)\.\s*(.*)', text)
    if m1:
        num = m1.group(1)
        rest = m1.group(2).strip()
    elif m2:
        num = m2.group(1)
        rest = m2.group(2).strip()
        
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Inches(0.25)
    
    if num:
        run_num = p.add_run(f"{num}. ")
        apply_text_formatting(run_num, bold=True)
        
    add_runs_to_paragraph(p, rest)

def process_table(doc, lines):
    """Parses markdown table lines and creates a formatted native Word Table."""
    if len(lines) < 2:
        return
        
    headers = [cell.strip() for cell in lines[0].split('|')[1:-1]]
    num_cols = len(headers)
    
    data_rows = []
    for line in lines[2:]:
        cells = [cell.strip() for cell in line.split('|')[1:-1]]
        while len(cells) < num_cols:
            cells.append("")
        data_rows.append(cells[:num_cols])
        
    num_rows = len(data_rows) + 1
    
    table = doc.add_table(rows=num_rows, cols=num_cols, style='Table Grid')
    
    hdr_cells = table.rows[0].cells
    for col_idx, text in enumerate(headers):
        hdr_cells[col_idx].text = ""
        p = hdr_cells[col_idx].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text)
        apply_text_formatting(run, size_pt=10, bold=True)
        set_cell_background(hdr_cells[col_idx], "D3D3D3")
        set_cell_margins(hdr_cells[col_idx], top=100, bottom=100, left=120, right=120)
        
    for row_idx, row_data in enumerate(data_rows):
        row_cells = table.rows[row_idx + 1].cells
        for col_idx, text in enumerate(row_data):
            row_cells[col_idx].text = ""
            p = row_cells[col_idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_runs_to_paragraph(p, text)
            for run in p.runs:
                run.font.size = Pt(10)
            set_cell_margins(row_cells[col_idx], top=80, bottom=80, left=120, right=120)
            
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)

def main():
    doc = Document()
    
    # Set page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    md_path = "SeedSec_Rwanda_Proposal.md"
    if not os.path.exists(md_path):
        print(f"Error: {md_path} not found.")
        return

    with open(md_path, "r", encoding="utf-8") as f:
        lines = f.readlines()

    in_cover_page = False
    in_references = False
    current_type = None
    buffer = []
    
    def flush():
        nonlocal current_type, buffer, in_references
        if not buffer:
            return
        block_text = " ".join(buffer)
        if current_type == 'paragraph':
            stripped_text = block_text.strip()
            is_fully_bold = stripped_text.startswith("**") and stripped_text.endswith("**")
            inner_text = stripped_text[2:-2].strip() if is_fully_bold else stripped_text
            if is_fully_bold and inner_text.upper() == "REFERENCES":
                in_references = True
            process_paragraph_block(doc, block_text, in_references)
        elif current_type == 'bullet_list':
            process_bullet_list_block(doc, buffer)
        elif current_type == 'numbered_list':
            process_numbered_list_block(doc, buffer)
        elif current_type == 'table':
            process_table(doc, buffer)
        buffer = []
        current_type = None

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        
        # Cover page logic
        if stripped == "# Cover Page":
            flush()
            in_cover_page = True
            i += 1
            continue
            
        if stripped.startswith("# ") and in_cover_page:
            flush()
            in_cover_page = False
            title = stripped[2:].strip()
            
            p_logo = doc.add_paragraph()
            p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_logo.paragraph_format.space_before = Pt(36)
            p_logo.paragraph_format.space_after = Pt(24)
            run_logo = p_logo.add_run("ALU\n")
            apply_text_formatting(run_logo, font_name="Arial", size_pt=28, bold=True)
            
            p_title = doc.add_paragraph()
            p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_title.paragraph_format.space_after = Pt(48)
            run_title = p_title.add_run(title)
            apply_text_formatting(run_title, size_pt=16, bold=True)
            
            p_degree = doc.add_paragraph()
            p_degree.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_degree.paragraph_format.space_after = Pt(96)
            run_deg = p_degree.add_run("BSc. in Software Engineering\nCapstone Research Proposal")
            apply_text_formatting(run_deg, size_pt=12, bold=True)
            
            details = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("---"):
                det_line = lines[i].strip()
                if det_line:
                    details.append(det_line)
                i += 1
            
            p_author = doc.add_paragraph()
            p_author.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_author.paragraph_format.space_after = Pt(24)
            for det in details:
                if ":" in det:
                    label, val = det.split(":", 1)
                    r_lbl = p_author.add_run(f"{label}:")
                    apply_text_formatting(r_lbl, size_pt=12, bold=True)
                    r_val = p_author.add_run(f"{val}\n")
                    apply_text_formatting(r_val, size_pt=12)
                else:
                    run_det = p_author.add_run(f"{det}\n")
                    apply_text_formatting(run_det, size_pt=12)
            
            doc.add_page_break()
            i += 1
            continue

        # If line is blank
        if stripped == "":
            flush()
            i += 1
            continue

        # Check block-initiating patterns
        is_h1 = stripped.startswith("# ")
        is_h2 = stripped.startswith("## ")
        is_h3 = stripped.startswith("### ")
        is_h4 = stripped.startswith("#### ")
        is_divider = stripped.startswith("---")
        is_table_row = stripped.startswith("|")
        is_image = re.match(r'^!\[(.*?)\]\((.*?)\)', stripped)
        
        is_bullet = stripped.startswith("* ") or stripped.startswith("- ")
        is_numbered = re.match(r'^\d+\.\s+', stripped) or re.match(r'^\*\*(\d+)\.\s*\*\*\s+', stripped)

        if is_divider:
            flush()
            doc.add_page_break()
            i += 1
            continue
            
        if is_image:
            flush()
            caption = is_image.group(1).strip()
            image_path = is_image.group(2).strip()
            if ")" in image_path:
                image_path = image_path.split(")")[0].strip()
            image_path = re.sub(r'\{.*\}', '', image_path).strip()
            image_path = image_path.replace("`", "").strip()
            
            full_img_path = image_path
            if not os.path.isabs(full_img_path):
                full_img_path = os.path.join(os.getcwd(), image_path)
            
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.paragraph_format.space_before = Pt(12)
            p_img.paragraph_format.space_after = Pt(6)
            
            if os.path.exists(full_img_path):
                p_img.add_run().add_picture(full_img_path, width=Inches(5.8))
            else:
                run_err = p_img.add_run(f"[Image Missing: {image_path}]")
                apply_text_formatting(run_err, bold=True)
                
            p_cap = doc.add_paragraph()
            p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_cap.paragraph_format.space_before = Pt(6)
            p_cap.paragraph_format.space_after = Pt(18)
            p_cap.paragraph_format.keep_with_next = True
            
            run_cap = p_cap.add_run(caption)
            apply_text_formatting(run_cap, size_pt=10, bold=True, italic=True)
            
            i += 1
            continue

        if is_h1 or is_h2 or is_h3 or is_h4:
            flush()
            if is_h1:
                h_text = stripped[2:].strip()
                if h_text.upper() == "REFERENCES":
                    in_references = True
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT if "CHAPTER" in h_text or "REFERENCES" in h_text else WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(24)
                p.paragraph_format.space_after = Pt(18)
                p.paragraph_format.keep_with_next = True
                run = p.add_run(h_text)
                apply_text_formatting(run, size_pt=16, bold=True)
            elif is_h2:
                h_text = stripped[3:].strip()
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.space_before = Pt(18)
                p.paragraph_format.space_after = Pt(12)
                p.paragraph_format.keep_with_next = True
                run = p.add_run(h_text)
                apply_text_formatting(run, size_pt=14, bold=True)
            elif is_h3:
                h_text = stripped[4:].strip()
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(6)
                p.paragraph_format.keep_with_next = True
                run = p.add_run(h_text)
                apply_text_formatting(run, size_pt=12, bold=True)
            elif is_h4:
                h_text = stripped[5:].strip()
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(6)
                p.paragraph_format.keep_with_next = True
                run = p.add_run(h_text)
                apply_text_formatting(run, size_pt=12, bold=True, italic=True)
            i += 1
            continue

        if is_table_row:
            if current_type != 'table':
                flush()
                current_type = 'table'
            buffer.append(stripped)
            i += 1
            continue

        if is_bullet:
            flush()
            current_type = 'bullet_list'
            buffer.append(stripped)
            i += 1
            continue

        if is_numbered:
            flush()
            current_type = 'numbered_list'
            buffer.append(stripped)
            i += 1
            continue

        # If it's normal text line
        if current_type is None:
            current_type = 'paragraph'
            buffer.append(stripped)
        elif current_type == 'paragraph':
            buffer.append(stripped)
        elif current_type in ('bullet_list', 'numbered_list'):
            buffer.append(stripped)
        elif current_type == 'table':
            flush()
            current_type = 'paragraph'
            buffer.append(stripped)

        i += 1

    flush()

    output_filename = "SeedSec_Rwanda_Proposal_Clean_Final.docx"
    doc.save(output_filename)
    print(f"Success: Generated {output_filename}")

if __name__ == "__main__":
    main()
